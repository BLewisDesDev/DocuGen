# File: create_template.py
"""
Script to create a sample Word template for testing.

Run this to create a test template with Jinja2 syntax.
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path

def create_care_plan_template():
    """Create a care plan template with Jinja2 syntax."""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('Caura CHSP CARE PLAN', 0)
    title.alignment = 1  # Center alignment
    
    # Client Information Section
    doc.add_heading('CLIENT INFORMATION', level=1)
    
    doc.add_paragraph('Client Name: {{ FirstName }} {{ LastName }}')
    doc.add_paragraph('Date of Birth: {{ DOB }}')
    doc.add_paragraph('Gender: {{ Gender }}')
    doc.add_paragraph('Address: {{ Address1 }} {{ Address2 }} {{ Suburb }} NSW {{ PostCode }}')
    
    # Plan Details Section
    doc.add_paragraph()  # Empty line
    doc.add_heading('PLAN DETAILS', level=1)
    
    doc.add_paragraph('Date of Plan: {{ DateOfPlan }}')
    doc.add_paragraph('Date of Next Review: {{ ReviewDate }}')
    
    # Services Provided Section
    doc.add_paragraph()  # Empty line
    doc.add_heading('SERVICES PROVIDED', level=1)
    
    # Home Maintenance
    hm_para = doc.add_paragraph('Home Maintenance:')
    hm_para.add_run('\n{% if Type == "HM" %}☑ Yes  ☐ No{% else %}☐ Yes  ☑ No{% endif %}')
    
    # Domestic Assistance
    da_para = doc.add_paragraph('Domestic Assistance:')
    da_para.add_run('\n{% if Type == "DA" %}☑ Yes  ☐ No{% else %}☐ Yes  ☑ No{% endif %}')
    
    # Personal Care
    pc_para = doc.add_paragraph('Personal Care:')
    pc_para.add_run('\n{% if Type == "PC" %}☑ Yes  ☐ No{% else %}☐ Yes  ☑ No{% endif %}')
    
    # Community Access
    ca_para = doc.add_paragraph('Community Access:')
    ca_para.add_run('\n{% if Type == "CA" %}☑ Yes  ☐ No{% else %}☐ Yes  ☑ No{% endif %}')
    
    # Staff Sign-off Section
    doc.add_paragraph()  # Empty line
    doc.add_heading('STAFF SIGN-OFF', level=1)
    
    doc.add_paragraph('Signed: _______________________________')
    doc.add_paragraph('Date: _________________________________')
    
    # Review Notes Section
    doc.add_paragraph()  # Empty line
    doc.add_heading('REVIEW NOTES', level=1)
    
    doc.add_paragraph('Date of Review: _________________')
    doc.add_paragraph()
    doc.add_paragraph('Changes to Care Plan:')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()
    doc.add_paragraph('Details:')
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    doc.add_paragraph()
    doc.add_paragraph('Next Review Date: _________________')
    doc.add_paragraph('Reviewed by: _____________________')
    
    # Save template
    template_dir = Path("templates")
    template_dir.mkdir(exist_ok=True)
    template_path = template_dir / "care_plan_template.docx"
    
    doc.save(template_path)
    print(f"✓ Created template: {template_path}")
    
    return template_path

if __name__ == "__main__":
    create_care_plan_template()