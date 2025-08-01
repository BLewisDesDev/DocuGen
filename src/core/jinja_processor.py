# File: src/templates/jinja_processor.py

import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Set
from docx import Document
from jinja2 import Environment, Template, TemplateError

class JinjaProcessor:
    """Handles Jinja2 template processing with Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.jinja_env = Environment()
    
    def extract_template_variables(self, template_path: Path) -> Set[str]:
        """Extract all Jinja2 variables from a Word document template."""
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            doc = Document(template_path)
            template_text = self._extract_all_text(doc)
            
            # Find all Jinja2 variables: {{ variable_name }}
            variable_pattern = r'\{\{\s*([^}]+)\s*\}\}'
            variables = set()
            
            for match in re.finditer(variable_pattern, template_text):
                var_expr = match.group(1).strip()
                # Extract base variable name (before filters or operations)
                base_var = var_expr.split('|')[0].split('.')[0].strip()
                variables.add(base_var)
            
            self.logger.debug(f"Found {len(variables)} variables in template: {variables}")
            return variables
            
        except Exception as e:
            self.logger.error(f"Failed to extract variables from template {template_path}: {str(e)}")
            raise
    
    def validate_template_syntax(self, template_path: Path) -> bool:
        """Validate Jinja2 syntax in the template."""
        
        try:
            doc = Document(template_path)
            template_text = self._extract_all_text(doc)
            
            # Try to parse the template with Jinja2
            template = self.jinja_env.from_string(template_text)
            
            self.logger.info(f"✅ Template syntax validation passed: {template_path}")
            return True
            
        except TemplateError as e:
            self.logger.error(f"Template syntax error in {template_path}: {str(e)}")
            raise
        except Exception as e:
            self.logger.error(f"Failed to validate template {template_path}: {str(e)}")
            raise
    
    def process_template(self, template_path: Path, data: Dict[str, Any]) -> Document:
        """Process a Word template with Jinja2 and return the rendered document."""
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            # Load the Word document
            doc = Document(template_path)
            
            # Process all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    rendered_text = self._render_text(original_text, data)
                    
                    if original_text != rendered_text:
                        # Clear the paragraph and add rendered text
                        paragraph.clear()
                        paragraph.add_run(rendered_text)
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                rendered_text = self._render_text(original_text, data)
                                
                                if original_text != rendered_text:
                                    paragraph.clear()
                                    paragraph.add_run(rendered_text)
            
            # Process headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    self._process_header_footer(section.header, data)
                
                # Footers
                if section.footer:
                    self._process_header_footer(section.footer, data)
            
            self.logger.debug(f"Template processed successfully: {template_path}")
            return doc
            
        except Exception as e:
            self.logger.error(f"Failed to process template {template_path}: {str(e)}")
            raise
    
    def _extract_all_text(self, doc: Document) -> str:
        """Extract all text from a Word document including headers, footers, and tables."""
        
        text_parts = []
        
        # Paragraphs
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text_parts.append(paragraph.text)
        
        # Headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    text_parts.append(paragraph.text)
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text_parts.append(paragraph.text)
        
        return '\n'.join(text_parts)
    
    def _render_text(self, text: str, data: Dict[str, Any]) -> str:
        """Render text with Jinja2 template engine."""
        
        try:
            # Create Jinja2 template from text
            template = self.jinja_env.from_string(text)
            
            # Render with data
            rendered = template.render(**data)
            
            return rendered
            
        except TemplateError as e:
            self.logger.warning(f"Template rendering error in text '{text[:50]}...': {str(e)}")
            return text  # Return original text on error
        
        except Exception as e:
            self.logger.warning(f"Unexpected error rendering text '{text[:50]}...': {str(e)}")
            return text
    
    def _process_header_footer(self, header_footer, data: Dict[str, Any]):
        """Process headers and footers."""
        
        for paragraph in header_footer.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                rendered_text = self._render_text(original_text, data)
                
                if original_text != rendered_text:
                    paragraph.clear()
                    paragraph.add_run(rendered_text)